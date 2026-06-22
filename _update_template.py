import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import re

doc = Document('PCPD MODELO.docx')
table = doc.tables[0]

def merge_paragraph_runs(paragraph):
    """Merge all runs in a paragraph into a single run, preserving first run's formatting."""
    if not paragraph.runs:
        return
    # Get full text
    full_text = ''.join(run.text for run in paragraph.runs)
    # Get formatting from first run
    first_run = paragraph.runs[0]
    fmt = {
        'bold': first_run.bold,
        'italic': first_run.italic,
        'size': first_run.font.size,
        'name': first_run.font.name,
    }
    # Remove all existing runs
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    # Add new single run
    new_run = paragraph.add_run(full_text)
    new_run.bold = fmt['bold']
    new_run.italic = fmt['italic']
    if fmt['size']:
        new_run.font.size = fmt['size']
    if fmt['name']:
        new_run.font.name = fmt['name']

# ─── Process specific rows ───
rows_to_process = {
    21: [  # Órgão cotista
        (r'do\(a\):\s*$', 'do(a): {{nr_nota_credito}}'),
    ],

    1: [  # Beneficiário
        (r'\(\s+\)\s*Militar', '{{cb_beneficiario_militar}} Militar'),
        (r'\(\s+\)\s*Servidor Civil', '{{cb_beneficiario_servidor_civil}} Servidor Civil'),
        (r'\(\s+\)\s*Colaborador Eventual', '{{cb_beneficiario_colaborador_eventual}} Colaborador Eventual'),
    ],
    13: [  # Alojamento
        (r'\(\s+\)\s*Sim\s+', '{{cb_alojamento_sim}} Sim '),
        (r'\(\s+\)\s*N[aã]o', '{{cb_alojamento_nao}} N\u00e3o'),
    ],
    15: [  # Veículo Oficial
        (r'\(\s+\)\s*Sim', '{{cb_veiculo_oficial_sim}} Sim'),
        (r'\(\s+\)\s*N[aã]o', '{{cb_veiculo_oficial_nao}} N\u00e3o'),
        (r'\(\s+\)\s*Em parte da viagem', '{{cb_veiculo_oficial_em_parte}} Em parte da viagem'),
    ],
    17: [  # Conexões (Art. 17, Portaria 290)
        (r'\(\s+\)\s*Sim\s+', '{{cb_conexoes_sim}} Sim '),
        (r'\(\s+\)\s*N[aã]o', '{{cb_conexoes_nao}} N\u00e3o'),
        (r'\(\s+\)\s*com Ve\u00edculo Oficial', '{{cb_conexoes_veiculo_oficial}} com Ve\u00edculo Oficial'),
        (r'\(\s+\)\s*com Veiculo Oficial', '{{cb_conexoes_veiculo_oficial}} com Veiculo Oficial'),
    ],
    19: [  # Categoria de transporte
        (r'\(\s+\)\s*rodovi[aá]rio', '{{cb_transporte_rodoviario}} rodovi\u00e1rio'),
        (r'\(\s+\)\s*a[ée]reo', '{{cb_transporte_aereo}} a\u00e9reo'),
        (r'\(\s+\)\s*ferrovi[aá]rio', '{{cb_transporte_ferroviario}} ferrovi\u00e1rio'),
        (r'\(\s+\)\s*aquavi[aá]rio', '{{cb_transporte_aquaviario}} aquavi\u00e1rio'),
        (r'\(\s+\)\s*Meios Pr[oó]prios', '{{cb_transporte_meios_proprios}} Meios Pr\u00f3prios'),
        (r'\(\s+\)\s*Viatura Oficial', '{{cb_transporte_viatura_oficial}} Viatura Oficial'),
    ],
}

for row_num, replacements in rows_to_process.items():
    cell = table.cell(row_num, 0)
    for p in cell.paragraphs:
        # Check if this paragraph has any of the relevant text
        full_text = ''.join(r.text for r in p.runs)
        has_relevant = False
        for pattern, _ in replacements:
            if re.search(pattern, full_text, re.IGNORECASE):
                has_relevant = True
                break
        if not has_relevant:
            continue
        
        # Merge all runs into one
        merge_paragraph_runs(p)
        
        # Apply replacements
        new_text = p.runs[0].text
        for pattern, replacement in replacements:
            new_text = re.sub(pattern, replacement, new_text)
        
        # Update text
        p.runs[0].text = new_text

temp_file = '__temp_modified.docx'
doc.save(temp_file)
os.replace(temp_file, 'PCPD MODELO.docx')

# Verify
doc2 = Document('PCPD MODELO.docx')
table2 = doc2.tables[0]
for ri in [1, 13, 15, 19]:
    cell = table2.cell(ri, 0)
    print(f'=== Row {ri} ===')
    for p in cell.paragraphs:
        print(repr(p.text[:350]))
    print()
